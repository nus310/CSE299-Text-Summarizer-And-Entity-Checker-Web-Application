from flask import Flask, redirect, url_for, request, render_template, jsonify, send_file
from nltk import text
from flask_bootstrap import Bootstrap

# NLP Packages
import spacy

nlp = spacy.load("en_core_web_sm")
from textblob import TextBlob


#other 
import json
from wordcloud import wordcloud
from flaskext.markdown import Markdown
import matplotlib.pyplot as plt
import time
from io import BytesIO
# from spacy import displacy 

#------------------ JINTU IMPORT STARTS----------------------------
from flask import Flask
from flask import render_template
from flask import request, url_for
import requests
import time


#sumy summary package
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer

#spacy summary package
import spacy
from spacy import displacy
nlp = spacy.load('en_core_web_sm')
from spacy.spacy_summarization import text_summarizer

#nltk summary package
import nltk
nltk.download('punkt')
nltk.download('stopwords')
from nltk.nltk_summarization import nltk_summarizer

# libraries for nltk and tdifd vectorizer
from nltk.nltk_tfidf_summarization import sent_tokenize, _create_frequency_matrix,  _create_tf_matrix, _create_documents_per_words,  _create_idf_matrix,  _create_tf_idf_matrix,  _score_sentences,  _find_average_score,  _generate_summary, word_drop
import math
from nltk import sent_tokenize, word_tokenize, PorterStemmer
from nltk.corpus import stopwords 

#gensim summary package
#from gensim.summarization import summarize

##################################  text classification start ###########################
#text classification Machine Learning package
import pandas as ps
import numpy as np
import sklearn
import joblib
import os
import matplotlib


# Load Our CountVectorizer
text_vectorizer = open("models/final_news_cv_vectorizer.pkl","rb")
text_countVectorizer = joblib.load(text_vectorizer)

# Load Our Models
def load_prediction_models(model_file):
	loaded_models = joblib.load(open(os.path.join(model_file),"rb"))
	return loaded_models

# show user the text type from the value calculated by the model
def get_keys(val, myDictionary):
    for key, value in myDictionary.items():
        if val == value:
            return key

##################################  text classification end ###########################

#pdf to text converter package
# from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
# from pdfminer.pdfpage import PDFPage
# from pdfminer.converter import TextConverter
# from pdfminer.layout import LAParams
# import io


# Web Scraping Pkg
from bs4 import BeautifulSoup
#from urllib.request import urlopen  

#------------------ JINTU IMPORT ENDS----------------------------

#Initialize app
app = Flask(__name__)
bootstrap = Bootstrap(app)
Markdown(app)


#------------------ JINTU FUNCTIONS STARTS----------------------------
# Function for Sumy Summarization
def sumy_summarizer(docx,summaryLine):
	parser = PlaintextParser.from_string(docx,Tokenizer("english"))
	lex_summarizer = LexRankSummarizer()
	summary = lex_summarizer(parser.document,summaryLine)
	summary_list = [str(sentence) for sentence in summary]
	result = ' '.join(summary_list)
	return result

# Fetch Text From Url
def get_text(raw_url):
    req_obj = requests.get(raw_url)
    fullText = req_obj.text
    soup = BeautifulSoup(fullText)
    #all_paras = soup.find_all("p")
    #wiki_text = ''
    #for para in all_paras:
    #    wiki_text += para.text 
    #return wiki_text
    fetched_text = ' '.join(map(lambda p:p.text,soup.find_all('p')))
    return fetched_text


# Reading Time
def readingTime(mytext):
    total_words = len([token.text for token in nlp(mytext)])
    #total_words = 5000
    estimatedTime = total_words/200.0
    return estimatedTime

#######################  summarized text save and download list start ###########################
# save file package
import base64
timestr = time.strftime("%Y%m%d-%H%M%S")

#template
HTML_WRAPPER = """<div style="overflow-x: auto; border: 1px solid #e6e9ef; border-radius: 0.25rem; padding: 1rem">{}</div>"""
file_name = 'yourdocument' + timestr + '.txt'


def writetofile(text,file_name):
	with open(os.path.join('downloads',file_name),'w') as f:
		f.write(text)
	return file_name


@app.route("/home")
def home():
    return render_template("homePage.html")

#summary from Text
@app.route("/analyze_summary", methods=['POST', 'GET'])
def analyze_summary():
    start = time.time()
    if request.method == 'POST':
        rawtext = request.form['rawtext']
        summaryLine = request.form['lineSlider']
        final_reading_time = readingTime(rawtext)
        final_summary = sumy_summarizer(rawtext,summaryLine)
        summary_reading_time = readingTime(final_summary)
        end = time.time()
        final_time = end-start
        lengthSummary=len(final_summary)
        lengthText=len(rawtext)
    return render_template("homePage.html",lengthText=lengthText,lengthSummary=lengthSummary, ctext = rawtext, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)

#summary from URL
@app.route("/analyze_url", methods=['POST', 'GET'])
def analyze_url():
    start = time.time()
    if request.method == 'POST':
         raw_url = request.form['raw_url']
         summaryLine = request.form['lineSlider']
         #rawtext = get_text(raw_url)
         rawtext = get_text(raw_url)
         final_reading_time = readingTime(rawtext)
         final_summary = sumy_summarizer(rawtext,summaryLine)
         summary_reading_time = readingTime(final_summary)
         end = time.time()
         final_time = end-start
         lengthSummary=len(final_summary)
         lengthText=len(rawtext)
         return render_template("homePage.html", lengthSummary=lengthSummary, lengthText=lengthText, ctext = rawtext, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)
import re
import string
#summary compare
@app.route("/compareSummary")
def compareSummary():
    return render_template("compareSummaries.html")




#compare from text
@app.route('/comparer',methods=['GET','POST'])
def comparer():
    start = time.time()
    if request.method == 'POST':
        rawtext = request.form['rawtext']
        rawtext1 = word_drop(rawtext)
        final_reading_time = readingTime(rawtext)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext1)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        # #final_summary_gensim = summarize(rawtext)
        final_summary_gensim = sumy_summarizer(rawtext1,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        sentences = sent_tokenize(rawtext1)
        total_documents = len(sentences)
        freq_matrix = _create_frequency_matrix(sentences)
        tf_matrix = _create_tf_matrix(freq_matrix)
        count_doc_per_words = _create_documents_per_words(freq_matrix)
        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
        sentence_scores = _score_sentences(tf_idf_matrix)
        threshold = _find_average_score(sentence_scores)
        final_summary_nltk = _generate_summary(sentences, sentence_scores, 1.1 * threshold)
        #final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)

#compare from url
@app.route('/comparerURL',methods=['GET','POST'])
def comparerURL():
    start = time.time()
    if request.method == 'POST':
        raw_url = request.form['raw_url']
        rawtext = get_text(raw_url)
        final_reading_time = readingTime(rawtext)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        # #final_summary_gensim = summarize(rawtext)
        final_summary_gensim = sumy_summarizer(rawtext,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)


##################################  text classification start ###########################
#Text type
@app.route("/textType")
def textType():
    return render_template("textClassifier.html")


# classify from text
@app.route('/classify',methods=['GET','POST'])
def classify():
    start = time.time()
    if request.method == 'POST':
        rawtext = request.form['rawtext']
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        #NaiveBayModel = open("models/newsclassifierNBmodel.pkl","wb")
        #clf = MultinomialNB()
        #predictionNB = clf.predict(vect_text)
        #finalResultNB = get_keys(predictionNB,prediction_labels)
        #joblib.dump(clf,NaiveBayModel)

        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)

# classify from URL
@app.route('/classifyURL',methods=['GET','POST'])
def classifyURL():
    start = time.time()
    if request.method == 'POST':
        raw_url = request.form['raw_url']
        rawtext = get_text(raw_url)
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)

##################################  text classification end ###########################



#########################  summarized text save and download list end ###########################
def make_downloadable(filename):
    readfile = open(os.path.join("downloads",filename)).read()
    b64 = base64.b64encode(readfile.encode()).decode()
    href = 'Download File File (right-click and save as <some_name>.txt)'.format(b64)
    return href

#summarized text save
@app.route("/save", methods=['POST', 'GET'])
def save():
    if request.method == 'POST':
        rawtext = request.form['rawtext']
        final_summary = sumy_summarizer(rawtext,3)
        file_to_download = writetofile(rawtext,file_name)
        info = "Saved Result As :: {}".format(file_name)
        #d_link = make_downloadable(file_to_download)
        #st.markdown(d_link,unsafe_allow_html=True)
        make_downloadable(file_to_download)
        return render_template('homePage.html',info=info)


#Download List
@app.route("/downloadListPage")
def downloadListPage():
    return render_template("downloadListPage.html")


#Download List download button
@app.route("/downloadList", methods=['POST', 'GET'])
def downloadList():
    files = os.listdir(os.path.join('downloads'))
    if request.method == 'POST':
        selectedDownloadList = request.form['List']
        file_to_download = selectedDownloadList
        final_summary= "final summary"
        writetofile(final_summary,selectedDownloadList)
        info = "File Name: {}".format(file_to_download)
        make_downloadable(file_to_download)
#    d_link = make_downloadable(file_to_download)
    #st.markdown(d_link,unsafe_allow_html=True)
    return render_template("downloadListPage.html", info=info)
#st.selectbox("Select File To Download",files)

#copy text
@app.route('/copy')
def copy():
    return render_template("homePage.html")

##################################  image to text conversion start ###########################
#using easyOCR
#import easyocr
#reader = easyocr.Reader(['en','en'])
#results = reader.readtext("a.jpg")
#text = ""
#for result in results:
#    text = text + result[1] + ""
#print(text)

#using pytesseract
import pytesseract as tess
# tess.pytesseract.tesseract_cmd = r'E:\Soft\Tesseract for cse299 project\tesseract.exe'
from PIL import Image
from flask import  redirect

app.config["IMAGE_UPLOADS"] = ".\images"

#image to text
@app.route('/picture', methods=['POST', 'GET'])
def picture():
    start = time.time()
    if request.method == 'POST' and 'photo' in request.files:
        image  = request.files['photo']
        #return redirect(request.url)
        image.save(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        
        # text extract from image
        img = Image.open(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        imageText = tess.image_to_string(img)
        #print(text5)

        #summary
        summaryLine = request.form['lineSlider']
        final_reading_time = readingTime(imageText)
        final_summary = sumy_summarizer(imageText,summaryLine)
        summary_reading_time = readingTime(final_summary)
        end = time.time()
        final_time = end-start
        lengthSummary=len(final_summary)
        lengthText=len(imageText)
        return render_template("homePage.html",lengthText=lengthText,lengthSummary=lengthSummary, ctext = imageText, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)


##################################  image to text conversion end ###########################
##################################  speech to text conversion start ###########################
from ibm_watson import SpeechToTextV1
from ibm_watson.websocket import RecognizeCallback, AudioSource 
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator

#API
apikey = '8cAmRWVWBYc2xFs58iel7LMn23SUQ-Jbe5ah7lTM16bI'
url = 'https://api.us-south.speech-to-text.watson.cloud.ibm.com/instances/1fd7f17a-8e9e-4436-8482-96400546110a'

# Setup Service
authenticator = IAMAuthenticator(apikey)
stt = SpeechToTextV1(authenticator=authenticator)
stt.set_service_url(url)

# speech to text
@app.route('/audio', methods=['POST', 'GET'])
def audio():
    start = time.time()
    if request.method == 'POST' and 'audio' in request.files:
        audioFile  = request.files['audio']
        #return redirect(request.url)
        audioFile.save(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename))
       
        #audio open
        #Perform conversion
        with open(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename), 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        audioText = res['results'][0]['alternatives'][0]['transcript']
        #print(text)

        #summary
        summaryLine = request.form['lineSlider']
        final_reading_time = readingTime(audioText)
        final_summary = sumy_summarizer(audioText,summaryLine)
        summary_reading_time = readingTime(final_summary)
        end = time.time()
        final_time = end-start
        lengthSummary=len(final_summary)
        lengthText=len(audioText)
        return render_template("homePage.html",lengthText=lengthText,lengthSummary=lengthSummary, ctext = audioText, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)


##################################  speech to text conversion end ###########################

##################################  video to text conversion start ###########################

#moviepy helps to convert video to audio file
from moviepy.editor import *

# video to text
@app.route('/video', methods=['POST', 'GET'])
def video():
    start = time.time()
    if request.method == 'POST' and 'video' in request.files:
        videoFile  = request.files['video']
        #return redirect(request.url)
        videoFile.save(os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename))
        
        # comvert video to audio file
        mp4_file = os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename)
        mp3_file = "images/audio.mp3"
        videoclip = VideoFileClip(mp4_file)
        audioclip = videoclip.audio
        audioclip.write_audiofile(mp3_file)
        

        with open("images/audio.mp3", 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        videoText = res['results'][0]['alternatives'][0]['transcript']
        

        #summary
        summaryLine = request.form['lineSlider']
        final_reading_time = readingTime(videoText)
        final_summary = sumy_summarizer(videoText,summaryLine)
        summary_reading_time = readingTime(final_summary)
        end = time.time()
        final_time = end-start
        lengthSummary=len(final_summary)
        lengthText=len(videoText)
        audioclip.close()
        videoclip.close()
        return render_template("homePage.html",lengthText=lengthText,lengthSummary=lengthSummary, ctext = videoText, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)

##################################  video to text conversion end ###########################


##################################  pdf to text conversion start ###########################
import PyPDF2

app.config["IMAGE_UPLOADS"] = ".\images"

# pdf to text
@app.route('/pdf', methods=['POST', 'GET'])
def pdf():
    start = time.time()
    if request.method == 'POST':
        if 'pdf' in request.files:
            pdfFile  = request.files['pdf']
            #return redirect(request.url)
            pdfFile.save(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))

            #pdf open
            a = PyPDF2.PdfFileReader(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))
            pdfText= " "
            r = a.getNumPages()
            for i in range(0,r):
                pdfText += a.getPage(i).extractText()
                #print(pdfText)

            #with open("text.txt", "w" ) as f:
            #    f.write(pdfText)

        #summary
        summaryLine = request.form['lineSlider']
        final_reading_time = readingTime(pdfText)
        final_summary = sumy_summarizer(pdfText,summaryLine)
        summary_reading_time = readingTime(final_summary)
        end = time.time()
        final_time = end-start
        lengthSummary=len(final_summary)
        lengthText=len(pdfText)
        return render_template("homePage.html",lengthText=lengthText,lengthSummary=lengthSummary, ctext = pdfText, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)

##################################  pdf to text conversion end ###########################
import docx

app.config["IMAGE_UPLOADS"] = ".\images"

# doc to text
@app.route('/doc', methods=['POST', 'GET'])
def doc():
    start = time.time()
    if request.method == 'POST':
        if 'doc' in request.files:
            docFile  = request.files['doc']
            #return redirect(request.url)
            docFile.save(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))

            docOpen = docx.Document(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))
            docText = docOpen.paragraphs[0].text

        #summary
        summaryLine = request.form['lineSlider']
        final_reading_time = readingTime(docText)
        final_summary = sumy_summarizer(docText,summaryLine)
        summary_reading_time = readingTime(final_summary)
        end = time.time()
        final_time = end-start
        lengthSummary=len(final_summary)
        lengthText=len(docText)
        return render_template("homePage.html",lengthText=lengthText,lengthSummary=lengthSummary, ctext = docText, final_summary = final_summary, final_time = final_time, final_reading_time = final_reading_time, summary_reading_time = summary_reading_time)

# ------------------------------------------------------






# ------------------SAIMA CODE STARTS--------------------------------


@app.route('/index_saima')
def index_saima():
	return render_template('index.html')

@app.route('/extract',methods=["GET","POST"])
def extract():
	if request.method == 'POST':
		raw_text = request.form['rawtext']
		docx = nlp(raw_text)
		html = displacy.render(docx,style="ent")
		html = html.replace("\n\n","\n")
		result = HTML_WRAPPER.format(html)

	return render_template('result.html',rawtext=raw_text,result=result)

@app.route('/previewer')
def previewer():
	return render_template('previewer.html')

@app.route('/preview',methods=["GET","POST"])
def preview():
	if request.method == 'POST':
		newtext = request.form['newtext']
		result = newtext

	return render_template('preview.html',newtext=newtext,result=result)












# ------------------SAIMA CODE ENDS--------------------------------






























@app.route("/")
def index():
    return render_template("Home.html")

@app.route("/wordinfo")
def wordinfo():
    return render_template("wordinfo.html")


@app.route("/analyze", methods = ["GET", "POST"])
def analyze():
    start = time.time()
    if request.method == "POST":
        rawtext = request.form["rawtext"]
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)

#--Text To SPEECH STARTS--
# from gtts import gTTS
# import os

# My_text_for_speech ="hello my name is nusrat islam"
# language = 'en'
# output = gTTS(text = My_text_for_speech, lang= language, slow = False)
# output.save("output.mp3")
# os.system("start output.mp3")


#--Text To SPEECH ENDS--




#API 
@app.route("/api")
def basic_api():
    return render_template("Api.html")


#API TOKENS
@app.route("/api/tokens/<string:mytext>", methods=['GET'])
def api_tokens(mytext):
    docx = nlp(mytext)
    #Token
    custom_tokens = [token.text for token in docx]
    return jsonify(mytext,custom_tokens)

# API Tokens & Lemma
@app.route("/api/lemma/<string:mytext>",methods=['GET'])
def api_lemma(mytext):
    docx = nlp(mytext.strip())
    # Tokens & Lemma
    custom_lemma = [('Token:{},Lemma:{}'.format(token.text, token.lemma_))for token in docx]
    return jsonify(mytext,custom_lemma)

#API NAMED ENTITY
#@app.route("/api/ner/<string:mytext>",methods=['GET'])
#def api_ner(mytext):
    #docx = nlp(mytext)
   # myNameEntities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
   # return jsonify(mytext,myNameEntities)

#API NAMED ENTITY
@app.route("/api/entities/<string:mytext>",methods=['GET'])
def api_entities(mytext):
    docx = nlp(mytext)
    myNameEntities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
    return jsonify(mytext,myNameEntities)


#API SENTIMENT 
@app.route("/api/sentiment/<string:mytext>",methods=['GET'])
def api_sentiment(mytext):
        blob = TextBlob(mytext)
        mySentiment = [mytext, blob.words, blob.sentiment]
        return jsonify(mySentiment)


# API ALL FEATURES OF WORD INFO & POS
# @app.route("/api/nlppify/<string:mytext>",methods=['GET'])
# def api_nlppify(mytext):
#    docx = nlp(mytext)
#     # Token
#    allData = ["Token:{}, Tag:{}, POS:{}, Dependency:{}, Lemma:{}, Shape:{}, Alpha:{}, IsStopword:{}".format(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
#    return jsonify(mytext,allData)


@app.route("/dashboard")
def dashboard():
    return render_template("dashboard.html")

@app.route("/Main_home")
def Main_home():
    return render_template("home.html")


@app.route("/images")
def imagescloud():
    return render_template("images.html")




#=============================new code test
#////////////////////// compare and text type from image start ////////////////////

######### compare from image start ##########
app.config["IMAGE_UPLOADS"] = ".\images"
#image to text and compare summaries
@app.route('/pictureCMP',methods=['GET','POST'])
def pictureCMP():
    start = time.time()
    if request.method == 'POST' and 'photo' in request.files:
        image  = request.files['photo']
        image.save(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        
        # text extract from image
        img = Image.open(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        imageText = tess.image_to_string(img)

        rawtext1 = imageText
        final_reading_time = readingTime(rawtext1)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext1)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        final_summary_gensim = sumy_summarizer(rawtext1,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        sentences = sent_tokenize(rawtext1)
        total_documents = len(sentences)
        freq_matrix = _create_frequency_matrix(sentences)
        tf_matrix = _create_tf_matrix(freq_matrix)
        count_doc_per_words = _create_documents_per_words(freq_matrix)
        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
        sentence_scores = _score_sentences(tf_idf_matrix)
        threshold = _find_average_score(sentence_scores)
        final_summary_nltk = _generate_summary(sentences, sentence_scores, threshold)
        #final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext1,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext1,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)




##### text classify from image ########
app.config["IMAGE_UPLOADS"] = ".\images"
#image to text and text classify
@app.route('/pictureTT',methods=['GET','POST'])
def pictureTT():
    start = time.time()
    if request.method == 'POST' and 'photo' in request.files:
        image  = request.files['photo']
        image.save(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        
        # text extract from image
        img = Image.open(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        imageText = tess.image_to_string(img)

        # summary code
        rawtext = imageText
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)


#////////////////////// compare from image end ////////////////////




#////////////////////// compare from audio start ////////////////////


app.config["IMAGE_UPLOADS"] = ".\images"
#audio to text and compare summaries
@app.route('/audioCMP',methods=['GET','POST'])
def audioCMP():
    start = time.time()
    if request.method == 'POST' and 'audio' in request.files:
        audioFile  = request.files['audio']
        audioFile.save(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename))
       
        #audio open
        #Perform conversion
        with open(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename), 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        audioText = res['results'][0]['alternatives'][0]['transcript']

        # compare code
        rawtext1 = audioText 
        final_reading_time = readingTime(rawtext1)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext1)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        final_summary_gensim = sumy_summarizer(rawtext1,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        sentences = sent_tokenize(rawtext1)
        total_documents = len(sentences)
        freq_matrix = _create_frequency_matrix(sentences)
        tf_matrix = _create_tf_matrix(freq_matrix)
        count_doc_per_words = _create_documents_per_words(freq_matrix)
        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
        sentence_scores = _score_sentences(tf_idf_matrix)
        threshold = _find_average_score(sentence_scores)
        final_summary_nltk = _generate_summary(sentences, sentence_scores, threshold)
        #final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext1,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext1,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)


######## text classify from audio ########
app.config["IMAGE_UPLOADS"] = ".\images"
#audio to text and text classify
@app.route('/audioTT',methods=['GET','POST'])
def audioTT():
    start = time.time()
    if request.method == 'POST' and 'audio' in request.files:
        audioFile  = request.files['audio']
        audioFile.save(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename))
       
        #audio open
        #Perform conversion
        with open(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename), 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        audioText = res['results'][0]['alternatives'][0]['transcript']

        # compare code
        rawtext = audioText
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)


#////////////////////// compare from audio end ////////////////////



#////////////////////// compare from video start ////////////////////

app.config["IMAGE_UPLOADS"] = ".\images"
#video to text and compare summaries
@app.route('/videoCMP',methods=['GET','POST'])
def videoCMP():
    start = time.time()
    if request.method == 'POST' and 'video' in request.files:
        videoFile  = request.files['video']
        #return redirect(request.url)
        videoFile.save(os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename))
        
        # comvert video to audio file
        mp4_file = os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename)
        mp3_file = "images/audio.mp3"
        videoclip = VideoFileClip(mp4_file)
        audioclip = videoclip.audio
        audioclip.write_audiofile(mp3_file)
        

        with open("images/audio.mp3", 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        videoText = res['results'][0]['alternatives'][0]['transcript']
    
        # compare code
        rawtext1 = videoText
        final_reading_time = readingTime(rawtext1)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext1)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        final_summary_gensim = sumy_summarizer(rawtext1,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        sentences = sent_tokenize(rawtext1)
        total_documents = len(sentences)
        freq_matrix = _create_frequency_matrix(sentences)
        tf_matrix = _create_tf_matrix(freq_matrix)
        count_doc_per_words = _create_documents_per_words(freq_matrix)
        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
        sentence_scores = _score_sentences(tf_idf_matrix)
        threshold = _find_average_score(sentence_scores)
        final_summary_nltk = _generate_summary(sentences, sentence_scores, threshold)
        #final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext1,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext1,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)


############ TEXT CLassify from video ################
app.config["IMAGE_UPLOADS"] = ".\images"
#video to text and text classify
@app.route('/videoTT',methods=['GET','POST'])
def videoTT():
    start = time.time()
    if request.method == 'POST' and 'video' in request.files:
        videoFile  = request.files['video']
        #return redirect(request.url)
        videoFile.save(os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename))
        
        # comvert video to audio file
        mp4_file = os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename)
        mp3_file = "images/audio.mp3"
        videoclip = VideoFileClip(mp4_file)
        audioclip = videoclip.audio
        audioclip.write_audiofile(mp3_file)
        

        with open("images/audio.mp3", 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        videoText = res['results'][0]['alternatives'][0]['transcript']
    
        # compare code
        rawtext = videoText
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)

#////////////////////// compare from video end ////////////////////




#////////////////////// compare from pdf start ////////////////////


app.config["IMAGE_UPLOADS"] = ".\images"
#pdf to text and compare summaries
@app.route('/pdfCMP',methods=['GET','POST'])
def pdfCMP():
    start = time.time()
    if request.method == 'POST':
        if 'pdf' in request.files:
            pdfFile  = request.files['pdf']
            #return redirect(request.url)
            pdfFile.save(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))

            #pdf open
            a = PyPDF2.PdfFileReader(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))
            pdfText= " "
            r = a.getNumPages()
            for i in range(0,r):
                pdfText += a.getPage(i).extractText()
    
        # compare code
        rawtext1 = pdfText
        final_reading_time = readingTime(rawtext1)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext1)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        final_summary_gensim = sumy_summarizer(rawtext1,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        sentences = sent_tokenize(rawtext1)
        total_documents = len(sentences)
        freq_matrix = _create_frequency_matrix(sentences)
        tf_matrix = _create_tf_matrix(freq_matrix)
        count_doc_per_words = _create_documents_per_words(freq_matrix)
        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
        sentence_scores = _score_sentences(tf_idf_matrix)
        threshold = _find_average_score(sentence_scores)
        final_summary_nltk = _generate_summary(sentences, sentence_scores, threshold)
        #final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext1,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext1,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)


########### text classify from pdf ################
app.config["IMAGE_UPLOADS"] = ".\images"
#pdf to text and text classifier
@app.route('/pdfTT',methods=['GET','POST'])
def pdfTT():
    start = time.time()
    if request.method == 'POST':
        if 'pdf' in request.files:
            pdfFile  = request.files['pdf']
            #return redirect(request.url)
            pdfFile.save(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))

            #pdf open
            a = PyPDF2.PdfFileReader(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))
            pdfText= " "
            r = a.getNumPages()
            for i in range(0,r):
                pdfText += a.getPage(i).extractText()
    
        # classify code
        rawtext = pdfText
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)


#////////////////////// compare from pdf end ////////////////////




#////////////////////// compare from doc start ////////////////////


app.config["IMAGE_UPLOADS"] = ".\images"
#doc to text and compare summaries
@app.route('/docCMP',methods=['GET','POST'])
def docCMP():
    start = time.time()
    if request.method == 'POST':
        if 'doc' in request.files:
            docFile  = request.files['doc']
            #return redirect(request.url)
            docFile.save(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))

            docOpen = docx.Document(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))
            docText = docOpen.paragraphs[0].text

        # compare code
        rawtext1 = docText
        final_reading_time = readingTime(rawtext1)
        #spacy Summarizer
        final_summary_spacy = text_summarizer(rawtext1)
        summary_reading_time = readingTime(final_summary_spacy)
        # Gensim Summarizer
        final_summary_gensim = sumy_summarizer(rawtext1,3)
        summary_reading_time_gensim = readingTime(final_summary_gensim)
		# NLTK Summarizer
        sentences = sent_tokenize(rawtext1)
        total_documents = len(sentences)
        freq_matrix = _create_frequency_matrix(sentences)
        tf_matrix = _create_tf_matrix(freq_matrix)
        count_doc_per_words = _create_documents_per_words(freq_matrix)
        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
        sentence_scores = _score_sentences(tf_idf_matrix)
        threshold = _find_average_score(sentence_scores)
        final_summary_nltk = _generate_summary(sentences, sentence_scores, threshold)
        #final_summary_nltk = nltk_summarizer(rawtext)
        summary_reading_time_nltk = readingTime(final_summary_nltk)
		# Sumy Summarizer
        final_summary_sumy = sumy_summarizer(rawtext1,3)
        summary_reading_time_sumy = readingTime(final_summary_sumy)
        
        end = time.time()
        final_time = end-start
        return render_template('compareSummaries.html',ctext=rawtext1,final_summary_spacy=final_summary_spacy,final_summary_gensim=final_summary_gensim,final_summary_nltk=final_summary_nltk,final_time=final_time,final_reading_time=final_reading_time,summary_reading_time=summary_reading_time,summary_reading_time_gensim=summary_reading_time_gensim,final_summary_sumy=final_summary_sumy,summary_reading_time_sumy=summary_reading_time_sumy,summary_reading_time_nltk=summary_reading_time_nltk)


###### text claasify from doc #######
app.config["IMAGE_UPLOADS"] = ".\images"
#doc to text and text classifier
@app.route('/docTT',methods=['GET','POST'])
def docTT():
    start = time.time()
    if request.method == 'POST':
        if 'doc' in request.files:
            docFile  = request.files['doc']
            #return redirect(request.url)
            docFile.save(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))

            docOpen = docx.Document(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))
            docText = docOpen.paragraphs[0].text

        # compare code
        rawtext = docText
        final_reading_time = readingTime(rawtext)
        prediction_labels = {'business':0,'tech':1,'sport':2,'health':3,'politics':4,'entertainment':5,'crime':6,'weather forecast':7}
        vect_text = text_countVectorizer.transform([rawtext]).toarray()
    
        # Logistic Regression Classifier
        predictorLR = load_prediction_models("models/newsclassifier_Logit_model.pkl")
        predictionLR = predictorLR.predict(vect_text)
        finalResultLR = get_keys(predictionLR,prediction_labels)
        # st.write(predictionLR)

        # Random Forest Classifier
        predictorRF = load_prediction_models("models/newsclassifier_RFOREST_model.pkl")
        predictionRF = predictorRF.predict(vect_text)
        finalResultRF = get_keys(predictionRF,prediction_labels)
        # st.write(predictionRF)

		# Decision Tree Classifier
        predictorDT = load_prediction_models("models/newsclassifier_CART_model.pkl")
        predictionDT = predictorDT.predict(vect_text)
        finalResultDT = get_keys(predictionDT,prediction_labels)
        # st.write(predictionDT)

		# Naive Bayes Classifier
        predictorNB = load_prediction_models("models/newsclassifier_NAIVEBAYES_model.pkl")
        predictionNB = predictorNB.predict(vect_text)
        finalResultNB = get_keys(predictionNB,prediction_labels)
        # st.write(predictionNB)
        
        end = time.time()
        final_time = end-start
        return render_template('textClassifier.html',ctext=rawtext,finalResultLR=finalResultLR,finalResultRF=finalResultRF,finalResultDT=finalResultDT,finalResultNB=finalResultNB,final_time=final_time,final_reading_time=final_reading_time)


#////////////////////// compare from doc end ////////////////////

#=========================rimi

#================url to text

@app.route('/url_word',methods=['GET','POST'])
def url_word():
    start = time.time()
    if request.method == 'POST':
        raw_url = request.form['raw_url']
        rawtext = get_text(raw_url)
  
       

        #--------
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)


      

#========= rimi image to text

app.config["IMAGE_UPLOADS"] = ".\images"
#image to text and compare summaries
@app.route('/pic_word',methods=['GET','POST'])
def pic_word():
    start = time.time()
    if request.method == 'POST' and 'photo' in request.files:
        image  = request.files['photo']
        image.save(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        
        # text extract from image
        img = Image.open(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        imageText = tess.image_to_string(img)

        rawtext = imageText

        #--------
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)

# audio to text

app.config["IMAGE_UPLOADS"] = ".\images"
#audio to text and compare summaries
@app.route('/audio_word',methods=['GET','POST'])
def audio_word():
    start = time.time()
    if request.method == 'POST' and 'audio' in request.files:
        audioFile  = request.files['audio']
        audioFile.save(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename))
       
        #audio open
        #Perform conversion
        with open(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename), 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        audioText = res['results'][0]['alternatives'][0]['transcript']

        # compare code
        rawtext = audioText 
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)


#

#////////////////////// compare from video start ////////////////////

app.config["IMAGE_UPLOADS"] = ".\images"
#video to text and compare summaries
@app.route('/video_word',methods=['GET','POST'])
def video_word():
    start = time.time()
    if request.method == 'POST' and 'video' in request.files:
        videoFile  = request.files['video']
        #return redirect(request.url)
        videoFile.save(os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename))
        
        # comvert video to audio file
        mp4_file = os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename)
        mp3_file = "images/audio.mp3"
        videoclip = VideoFileClip(mp4_file)
        audioclip = videoclip.audio
        audioclip.write_audiofile(mp3_file)
        

        with open("images/audio.mp3", 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        videoText = res['results'][0]['alternatives'][0]['transcript']
    
        # compare code
        rawtext = videoText
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)


#

#////////////////////// compare from pdf start ////////////////////


app.config["IMAGE_UPLOADS"] = ".\images"
#pdf to text and compare summaries
@app.route('/pdf_word',methods=['GET','POST'])
def pdf_word():
    start = time.time()
    if request.method == 'POST':
        if 'pdf' in request.files:
            pdfFile  = request.files['pdf']
            #return redirect(request.url)
            pdfFile.save(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))

            #pdf open
            a = PyPDF2.PdfFileReader(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))
            pdfText= " "
            r = a.getNumPages()
            for i in range(0,r):
                pdfText += a.getPage(i).extractText()
    
        # compare code
        rawtext = pdfText
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)

        
       

#
#////////////////////// compare from doc start ////////////////////


app.config["IMAGE_UPLOADS"] = ".\images"
#doc to text and compare summaries
@app.route('/doc_word',methods=['GET','POST'])
def doc_word():
    start = time.time()
    if request.method == 'POST':
        if 'doc' in request.files:
            docFile  = request.files['doc']
            #return redirect(request.url)
            docFile.save(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))

            docOpen = docx.Document(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))
            docText = docOpen.paragraphs[0].text

        # compare code
        rawtext = docText
        #--------
        text_length = len(rawtext) # for the time of text analysis
        #Analysis
        docx = nlp(rawtext)
        #Token
        custom_tokens = [token.text for token in docx]
        #WordInfo
        custom_wordInfo = [(token.text, token.lemma_, token.shape_,token.is_alpha,token.is_stop) for token in docx]
        #POS
        custom_pos = [(token.text, token.tag_, token.pos_, token.dep_) for token in docx]
        #Entities
        custom_entities = [(entity.text, entity.label_ , str(spacy.explain(entity.label_))) for entity in docx.ents]
        #Sentiments
        blob = TextBlob(rawtext)
        blob_sentiment = blob.sentiment.polarity
        blob_subjectivity = blob.sentiment.subjectivity
        result_json = json.dumps(custom_wordInfo, sort_keys = False, indent= 2)
        
        end = time.time() # for the time of text analysis
        final_time = end - start # for the time of text analysis
        # doc = nlp(rawtext)
        # displacy = displacy.render(doc, style = 'dep',jupyter = True , options = { 'distance' : 110 } )

        return render_template( "wordInfo.html",ctext = rawtext, text_length = text_length, final_time = final_time, custom_tokens = custom_tokens, custom_wordInfo = custom_wordInfo, result_json = result_json, custom_pos = custom_pos, custom_entities = custom_entities,blob_subjectivity = blob_subjectivity, blob_sentiment = blob_sentiment)

        





#====================SAIMA CODE




# @app.route('/extract',methods=["GET","POST"])
# def extract():
# 	if request.method == 'POST':
# 		raw_text = request.form['rawtext']
# 		docx = nlp(raw_text)
# 		html = displacy.render(docx,style="ent")
# 		html = html.replace("\n\n","\n")
# 		result = HTML_WRAPPER.format(html)

# 	return render_template('result.html',rawtext=raw_text,result=result)

# @app.route('/previewer')
# def previewer():
# 	return render_template('previewer.html')

# app.route('/preview',methods=["GET","POST"])
# def preview():
# 	if request.method == 'POST':
# 		newtext @= request.form['newtext']
# 		result = newtext

# 	return render_template('preview.html',newtext=newtext,result=result)

# url to text

#pdf to text and compare summaries
@app.route('/url_en',methods=['GET','POST'])
def url_en():
    start = time.time()
    if request.method == 'POST':
        raw_url = request.form['raw_url']
        raw_text = get_text(raw_url)
    
        # compare code
    
        docx = nlp(raw_text)
        html = displacy.render(docx,style="ent")
        html = html.replace("\n\n","\n")
        result = HTML_WRAPPER.format(html)
    return render_template('result.html',rawtext=raw_text,result=result)





# saima image to text


app.config["IMAGE_UPLOADS"] = ".\images"
#image to text and compare summaries
@app.route('/pic_en',methods=['GET','POST'])
def pic_en():
    start = time.time()
    if request.method == 'POST' and 'photo' in request.files:
        image  = request.files['photo']
        image.save(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        
        # text extract from image
        img = Image.open(os.path.join(app.config["IMAGE_UPLOADS"], image.filename))
        imageText = tess.image_to_string(img)

        rawtext = imageText
        docx = nlp(rawtext)
        html = displacy.render(docx,style="ent")
        html = html.replace("\n\n","\n")
        result = HTML_WRAPPER.format(html)
        return render_template('result.html',rawtext=rawtext,result=result)


# saima audio to text



app.config["IMAGE_UPLOADS"] = ".\images"
#audio to text and compare summaries
@app.route('/audio_en',methods=['GET','POST'])
def audio_en():
    start = time.time()
    if request.method == 'POST' and 'audio' in request.files:
        audioFile  = request.files['audio']
        audioFile.save(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename))
       
        #audio open
        #Perform conversion
        with open(os.path.join(app.config["IMAGE_UPLOADS"], audioFile.filename), 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        audioText = res['results'][0]['alternatives'][0]['transcript']

        # compare code
        raw_text = audioText 
        docx = nlp(raw_text)
        html = displacy.render(docx,style="ent")
        html = html.replace("\n\n","\n")
        result = HTML_WRAPPER.format(html)
    return render_template('result.html',rawtext=raw_text,result=result)

#

#////////////////////// compare from video start ////////////////////

app.config["IMAGE_UPLOADS"] = ".\images"
#video to text and compare summaries
@app.route('/video_en',methods=['GET','POST'])
def video_en():
    start = time.time()
    if request.method == 'POST' and 'video' in request.files:
        videoFile  = request.files['video']
        #return redirect(request.url)
        videoFile.save(os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename))
        
        # comvert video to audio file
        mp4_file = os.path.join(app.config["IMAGE_UPLOADS"], videoFile.filename)
        mp3_file = "images/audio.mp3"
        videoclip = VideoFileClip(mp4_file)
        audioclip = videoclip.audio
        audioclip.write_audiofile(mp3_file)
        

        with open("images/audio.mp3", 'rb') as f:
            res = stt.recognize(audio=f, content_type='audio/mp3', model='en-US_NarrowbandModel', continuous=True).get_result()
        videoText = res['results'][0]['alternatives'][0]['transcript']
    
        # compare code
        raw_text = videoText
        docx = nlp(raw_text)
        html = displacy.render(docx,style="ent")
        html = html.replace("\n\n","\n")
        result = HTML_WRAPPER.format(html)
    return render_template('result.html',rawtext=raw_text,result=result)



    #
    #////////////////////// compare from pdf start ////////////////////


app.config["IMAGE_UPLOADS"] = ".\images"
#pdf to text and compare summaries
@app.route('/pdf_en',methods=['GET','POST'])
def pdf_en():
    start = time.time()
    if request.method == 'POST':
        if 'pdf' in request.files:
            pdfFile  = request.files['pdf']
            #return redirect(request.url)
            pdfFile.save(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))

            #pdf open
            a = PyPDF2.PdfFileReader(os.path.join(app.config["IMAGE_UPLOADS"], pdfFile.filename))
            pdfText= " "
            r = a.getNumPages()
            for i in range(0,r):
                pdfText += a.getPage(i).extractText()
    
        # compare code
        raw_text = pdfText
        docx = nlp(raw_text)
        html = displacy.render(docx,style="ent")
        html = html.replace("\n\n","\n")
        result = HTML_WRAPPER.format(html)
    return render_template('result.html',rawtext=raw_text,result=result)

    #    
app.config["IMAGE_UPLOADS"] = ".\images"
#doc to text and compare summaries
@app.route('/doc_en',methods=['GET','POST'])
def doc_en():
    start = time.time()
    if request.method == 'POST':
        if 'doc' in request.files:
            docFile  = request.files['doc']
            #return redirect(request.url)
            docFile.save(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))

            docOpen = docx.Document(os.path.join(app.config["IMAGE_UPLOADS"], docFile.filename))
            docText = docOpen.paragraphs[0].text

        # compare code
        raw_text = docText
        docx = nlp(raw_text)
        html = displacy.render(docx,style="ent")
        html = html.replace("\n\n","\n")
        result = HTML_WRAPPER.format(html)
    return render_template('result.html',rawtext=raw_text,result=result)

        

        
        
        








if __name__ == "__main__":
    app.run(debug = True)